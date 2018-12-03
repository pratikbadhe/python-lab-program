# -*- coding: utf-8 -*-
"""
Created on Mon Nov 19 12:29:17 2018

@author: angish
"""
#angishbhandwalkar m-03

from easygui import *
import sys
import Tkinter
import tkMessageBox
from docx.shared import Inches

msgbox("Bonjour Monsiuer/Mademosielle")
while 1:
 a = choicebox("Apply for","E-Seva Kendra",choices=["Passport","Aadhar Card"])
 if a == "Passport":
  a0 = ccbox("You have opted for Passport Services of India")
  fieldNames = ["Enter applicant's full name(in caps)","Birth Date","State","City","District","Enter Pan card number 'if available' or enter '-' if not available","Enter Aadhar card number","Enter Voter id 'if available' or enter '-' if not available"]
  fieldValues = []
  fieldValues = multenterbox("Enter Details","Personal Info",fieldNames)
  a22= fileopenbox("upload your photo")
  a6 = choicebox("Gender",choices= ["Male","Female","Other"])
  a7 = choicebox("Are you married?","Marital status",choices = ["Yes","No"])
  a11 = choicebox("Employment type",choices = ["Student","Business","Government sector","Employee","Housewife","none"])
  a12 = choicebox("Is either of your parent (In case of minor i.e age below 18) a government servant",choices = ["Yes","No"])
  a13 = choicebox("Education qualification",choices =["<10","Matric","10+2","10+2[+]Graduation","10+2[+]Graduation[+]Post Graduation","Diploma"])
  #a14 = choicebox("Provide atleast one's details",choices = ["Father","Mother","Legal guardian"])
  a14 = choicebox("Is any of your family member serving in Indian armed forces",choices = ["Yes","No"])
  if a14 == "Yes": 
   a15 = enterbox("Enter your relation with the person who is in Army")
  else: 
   a15 = "No" 
  c1 = msgbox("I solemly swear that the above mentioned information is true. \n\t\t Yours faithfully, \n ______")
  from docx import Document
  document=Document()
  document.add_picture(a22, width=Inches(1.0))
  #paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')
  document.save('Passport.docx')
  f = open('Passport.docx', 'rb')
  document = Document(f)
  paragraph=document.add_paragraph("name="+ fieldValues[0]+"\nBirth date="+ fieldValues[1]+"\nstate="+ fieldValues[2]+"\nCity="+ fieldValues[3]+"\nDistrict="+ fieldValues[4]+"\nPancard Number="+fieldValues[5]+"\nAdhaar Number="+ fieldValues[6]+"\nVoter ID="+ fieldValues[7]+"\nGender=" + a6+"\nMartial Status"+ a7 +"\nEmployment=" +a11+ "\nAre parents government servant =" + a12+"\nEducation Qualification =" + a13+"\nrelation with person if serving in Indian army ="+ a15)
  document.save('Passport.docx')
  sys.exit(0)
 elif a == "Aadhar Card":
  b = ccbox("You have opted for UIDAI (Unique Identification Authority of India) service")
  fieldNames = ["Enter applicant's full name(in caps)","Birth Date","State","City","District"]
  fieldValues = []
  fieldValues = multenterbox("Enter Details","Personal Info",fieldNames)
  b6 = choicebox("Gender",choices= ["Male","Female","Other"])
  b7 = choicebox("Are you married?","Marital status",choices = ["Yes","No"])
  a22=fileopenbox("upload your photo")
  a24=fileopenbox("upload address proof")
  c = msgbox("I confirm that I have been residing in India for at least 182 days in the preceding 12 months & information (including biometrics) provided by me to the UIDAI is my own and is true, correct and accurate. I am aware that my information including biometrics  will be used for generation of Aadhaar and authentication. I understand that my identity  information (except core biometric) may be provided to an agency only with my consent  during authentication or as per the provisions of the Aadhaar Act. I have a right to access my  identity information (except core biometrics) following the procedure laid down by UIDAI.  Verifier’s Stamp and Signature: (Verifier must put his/her Name, if stamp is not available)  Applicant’s signature/Thumbprint  ---------------------------------------------------------------------------------------------------------------------------  ---------------------------------------------- To be filled by the Enrolment Agency only : Date & time  of Enrolment: ---------------------------------------------------- “(Note: Incase of minor, the signature  will be done by parent/guardian. Incase of incapacitated person, the signature will be done by Legal Guardian of Incapacitated Person)")
  from docx import Document
  document=Document()
  document.add_picture(a22,width=Inches(1.0))
  
  #paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')
  document.save('Aadhar.docx')
  f = open('Aadhar.docx', 'rb')
  document = Document(f)
  paragraph=document.add_paragraph("name ="+ fieldValues[0]+"\nBirth date="+ fieldValues[1]+"\nstate="+ fieldValues[2]+"\nCity="+ fieldValues[3]+"\nDistrict="+ fieldValues[4]+"\nGender=" + b6+"\nMartial Status"+ b7  )
  document.add_picture(a24,width=Inches(4.0),height=Inches(4.0))
  document.save('Aadhar.docx')
  sys.exit(0)
 else:
  sys.exit(0)
